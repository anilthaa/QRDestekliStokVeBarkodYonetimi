#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""QR Destekli Stok ve Barkod Yönetimi — Özet Teknik Dokümantasyon üretici."""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "QRDestekliStokVeBarkodYonetimi_TeknikDokumantasyon.docx"
DOC_DATE = date.today().strftime("%d.%m.%Y")
VERSION = "2.3"
SQL_SCHEMA = Path(__file__).resolve().parent / "sql" / "01_create_tables.sql"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_sql_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    r = p.add_run(code.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(9)


TABLE_DESCRIPTIONS: dict[str, str] = {
    "KullaniciTip": "Rol grupları. Varsayilan=true olan tip yeni kayıt (/register) kullanıcılarına atanır.",
    "Kullanicilar": "Sistem kullanıcıları. Sifre alanı PBKDF2 hash saklar.",
    "Form": "Menü ve sayfa tanımları (RBAC). IsMenu=false kayıtlar yetki alt sayfaları içindir.",
    "KullaniciDetay": "Kullanıcıya özel sayfa yetkisi. Yetki: 0=Gizli, 1=Okuma, 2=Yazma.",
    "KullaniciTipDetay": "Kullanıcı tipine (role) form bazlı yetki ataması.",
    "Birim": "Ölçü birimleri (adet, kg, litre vb.).",
    "Kategoriler": "Ürün kategorileri.",
    "Urunler": "Ürün kartları ve güncel stok miktarı.",
    "StokHareketleri": "Stok giriş/çıkış kayıtları. HareketTipi: 1=Giriş, 2=Çıkış.",
}

TABLE_DOC_ORDER = [
    "KullaniciTip",
    "Kullanicilar",
    "Form",
    "KullaniciDetay",
    "KullaniciTipDetay",
    "Birim",
    "Kategoriler",
    "Urunler",
    "StokHareketleri",
]


def parse_table_ddls(sql_path: Path) -> dict[str, str]:
    text = sql_path.read_text(encoding="utf-8")
    blocks: dict[str, str] = {}
    current_name: str | None = None
    current_lines: list[str] = []

    for line in text.splitlines():
        if line.startswith("-- ── "):
            if current_name and current_lines:
                blocks[current_name] = "\n".join(current_lines).strip()
            marker = line.removeprefix("-- ── ").strip()
            current_name = marker.split()[0] if marker else None
            current_lines = []
            continue
        if line.startswith("BEGIN;") or line.startswith("COMMIT;"):
            continue
        if line.startswith("--") and not line.startswith("-- HareketTipi"):
            if current_name is None:
                continue
        if current_name is not None:
            current_lines.append(line)

    if current_name and current_lines:
        blocks[current_name] = "\n".join(current_lines).strip()

    return blocks


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("QR DESTEKLİ STOK VE BARKOD YÖNETİMİ")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Teknik Dokümantasyon (Özet)").font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Sürüm: {VERSION}  |  Tarih: {DOC_DATE}\n")
    meta.add_run(".NET 8 · Blazor Server · Radzen · PostgreSQL · Dapper")


def build_content(doc: Document) -> None:
    # 1. Proje amacı
    add_heading(doc, "1. Proje Amacı", 1)
    add_para(
        doc,
        "QR Destekli Stok ve Barkod Yönetimi; işletmelerin ürün stoklarını web tabanlı bir panel "
        "üzerinden yönetmesini, barkod/QR kod ile stok giriş-çıkış kaydı yapmasını ve kullanıcı "
        "rollerine göre yetkilendirilmiş erişim sağlamasını amaçlayan bir stok yönetim uygulamasıdır.",
    )
    add_para(
        doc,
        "Temel işlevler: ürün/kategori/birim tanımları, stok hareketleri, dashboard özeti, "
        "kullanıcı ve yetki yönetimi, profil ve şifre sıfırlama, liste dışa aktarma (Excel/CSV) "
        "ve QR kod ile kamuya açık ürün detay sayfası.",
    )

    # 2. Veritabanı tabloları
    add_heading(doc, "2. Veritabanı Tablo Yapısı", 1)
    add_para(
        doc,
        "PostgreSQL üzerinde soft delete (DelUser/DelDate) kullanılır. Ortak alanlar: ID, CreUser, CreDate, "
        "UpdUser, UpdDate, DelUser, DelDate. Her tablonun CREATE TABLE kodu aşağıdadır; toplu çalıştırma için "
        "Scripts/sql/01_create_tables.sql dosyası da mevcuttur.",
    )

    table_ddls = parse_table_ddls(SQL_SCHEMA)
    for table_name in TABLE_DOC_ORDER:
        add_heading(doc, table_name, 2)
        add_para(doc, TABLE_DESCRIPTIONS.get(table_name, ""))
        ddl = table_ddls.get(table_name)
        if ddl:
            add_sql_code(doc, ddl)
        else:
            add_para(doc, f"(DDL bulunamadı: {SQL_SCHEMA})")
        doc.add_paragraph()

    # 3. Sayfalar
    add_heading(doc, "3. Sayfalar", 1)
    add_table(
        doc,
        ["Rota", "Sayfa", "Amaç"],
        [
            ["/", "Ana Sayfa (Dashboard)", "Özet istatistikler, kritik stok listesi, son hareketler; stok sayfalarına kısayol kartları"],
            ["/login", "Giriş", "Kullanıcı oturum açma"],
            ["/register", "Kayıt Ol", "Yeni kullanıcı kaydı"],
            ["/sifremi-unuttum", "Şifremi Unuttum", "E-posta ile şifre sıfırlama bağlantısı talebi"],
            ["/sifre-sifirla", "Şifre Sıfırla", "Token ile yeni parola belirleme"],
            ["/hesap-onay", "Hesap Onay", "E-posta doğrulama token'ı ile hesap aktivasyonu"],
            ["/profil", "Profil", "Kullanıcı bilgisi düzenleme, profil fotoğrafı yükleme/kırpma, şifre değiştirme"],
            ["/erisim-engeli", "Erişim Engeli", "Yetkisiz sayfa erişiminde gösterilen uyarı"],
            ["/birimler", "Birim Listesi", "Ölçü birimlerini listeleme, ekleme, düzenleme, silme"],
            ["/kategoriler", "Kategori Listesi", "Ürün kategorilerini yönetme"],
            ["/urunler", "Ürün Listesi", "Ürün CRUD, barkod/QR görüntüleme, dışa aktarma"],
            ["/urun/{BarkodNo}", "Ürün Detay (Genel)", "QR kod ile erişilen kamuya açık ürün bilgi sayfası"],
            ["/kullanicilar", "Kullanıcı Listesi", "Kullanıcı yönetimi, hesap onay maili gönderme"],
            ["/kullanici-tipler", "Kullanıcı Tipi Listesi", "Rol/tip tanımları yönetimi"],
            ["/kullanici-yetki/{id}", "Kullanıcı Yetki", "Tek kullanıcıya form bazlı yetki atama"],
            ["/kullanici-tip-yetki/{id}", "Kullanıcı Tipi Yetki", "Role form bazlı yetki atama"],
            ["/stok-islemleri", "Stok İşlemleri", "Stok giriş veya çıkış sayfasına yönlendirme merkezi"],
            ["/stok-giris", "Stok Giriş", "Barkod/manuel ürün seçimi ile depoya giriş kaydı"],
            ["/stok-cikis", "Stok Çıkış", "Barkod/manuel ürün seçimi ile depodan çıkış kaydı"],
            ["/stok-hareketleri", "Stok Hareketleri", "Tüm giriş/çıkış geçmişi; tarih ve tip filtreleme"],
        ],
    )
    add_para(doc, "Diyalog bileşenleri (ayrı rota yok): BirimTnm, KategoriTnm, UrunTnm, KullaniciTnm, KullaniciTipTnm — ilgili listelerden açılan ekleme/düzenleme formları. ProfilResimKirpma — profil fotoğrafı kırpma diyalogu.")

    # 4. Servisler
    add_heading(doc, "4. Servisler", 1)
    add_table(
        doc,
        ["Servis", "Amaç"],
        [
            ["DataService", "Tüm CRUD işlemleri ve SQL sorguları; stok hareketi ile ürün stok güncelleme"],
            ["DBClass", "PostgreSQL bağlantı yönetimi"],
            ["AuthStateService", "Oturum açmış kullanıcı bilgisini bellekte tutma"],
            ["KullaniciCookieSignInService", "Cookie tabanlı oturum açma/kapama"],
            ["YetkiService", "Form bazlı RBAC; menü ve sayfa erişim kontrolü"],
            ["JwtService", "JWT token üretim/doğrulama (yardımcı)"],
            ["PasswordHasher", "Parola hashleme ve doğrulama (PBKDF2)"],
            ["QrService", "Ürün için QR kod görüntüsü üretimi"],
            ["ExportService", "Liste verilerini Excel/CSV olarak dışa aktarma"],
            ["EmailService", "SMTP üzerinden e-posta gönderimi"],
            ["EpostaDogrulamaService", "Hesap onay ve bilgilendirme e-posta şablonları"],
            ["SifreSifirlamaTokenService", "Şifre sıfırlama token üretimi ve doğrulama"],
            ["SifreDegistirDogrulamaService", "Parola güçlülük kuralları doğrulama"],
            ["KullaniciHesapOnayTokenService", "Hesap onay token üretimi ve doğrulama"],
            ["ProfilResimIslemci", "Profil fotoğrafı boyutlandırma ve kaydetme"],
            ["ProfilResimKirpState", "Kırpma diyalogundan profil sayfasına görsel aktarımı (scoped)"],
            ["UrunMetinDogrulama", "Ürün formu karakter sınırı kontrolü"],
            ["UrunResimYardimci", "Ürün görsel yolu ve önizleme yardımcıları"],
            ["GridSiraYardimci", "Grid satır sıra numarası hesaplama"],
            ["BarkodBasariSesiYardimci", "Barkod okuma başarı sesi tetikleme"],
            ["RequestBaseUrlHelper", "Dinamik uygulama base URL (e-posta linkleri)"],
        ],
    )

    # 5. Modeller
    add_heading(doc, "5. Modeller", 1)
    add_table(
        doc,
        ["Model", "Karşılık / Amaç"],
        [
            ["ItemBase", "Tüm varlıkların ortak alanları (ID, audit, soft delete)"],
            ["ItemUrun", "Urunler tablosu"],
            ["ItemKategori", "Kategoriler tablosu"],
            ["ItemBirim", "Birim tablosu"],
            ["ItemStokHareketleri", "StokHareketleri tablosu (+ JOIN görüntüleme alanları)"],
            ["ItemKullanicilar", "Kullanicilar tablosu"],
            ["ItemKullaniciTip", "KullaniciTip tablosu"],
            ["ItemForm", "Form tablosu (menü/sayfa)"],
            ["ItemKullaniciDetay", "KullaniciDetay tablosu"],
            ["ItemKullaniciTipDetay", "KullaniciTipDetay tablosu"],
            ["DataResult<T>", "Servis işlem sonucu sarmalayıcı (başarı/hata/mesaj)"],
            ["AuthResult", "Kimlik doğrulama işlem sonucu"],
            ["DashboardOzet", "Ana sayfa özet istatistik DTO"],
            ["YetkiTipi", "Yetki seviye sabitleri (Gizli, Okuma, Yazma)"],
            ["PermissionLevelOption", "Yetki seçim listesi UI modeli"],
            ["UrunAlanSinirlari", "Ürün formu alan uzunluk limitleri"],
            ["BarkodOkumaBildirimi", "Barkod okuma UI bildirim durumu"],
        ],
    )

    # 6. Sıfırdan kurulum
    add_heading(doc, "6. Sıfırdan Kurulum ve Devreye Alma", 1)
    add_para(
        doc,
        "Proje başka bir ortama sıfır olarak devredildiğinde aşağıdaki adımlar izlenir. "
        "Veritabanı kurulumu için repodaki Scripts/sql/ klasöründeki PostgreSQL scriptleri kullanılır.",
    )

    add_heading(doc, "6.1 Gereksinimler", 2)
    add_bullets(
        doc,
        [
            ".NET 8 SDK (geliştirme) veya .NET 8 ASP.NET Core Runtime (sunucu)",
            "PostgreSQL 14 veya üzeri",
            "Windows/Linux sunucu veya IIS + ASP.NET Core Hosting Bundle (Windows IIS için)",
            "SMTP sunucusu (isteğe bağlı; şifre sıfırlama ve hesap onay e-postası için)",
        ],
    )

    add_heading(doc, "6.2 Kurulum Adımları", 2)
    steps = [
        (
            "1. Kaynak kodunu sunucuya kopyalayın",
            "Projeyi zip/repo olarak alın. Üretimde `dotnet publish -c Release -o ./publish` ile yayın paketi oluşturulabilir.",
        ),
        (
            "2. PostgreSQL veritabanını hazırlayın",
            "Boş bir veritabanı oluşturun (ör. StokVeBarkodYonetimiDB). Scripts/sql/00_kurulum.sql dosyasını "
            "psql ile çalıştırın (veya sırayla 01_create_tables.sql + 02_seed_data.sql). "
            "02_seed_data.sql Form menü kayıtlarını ve varsayılan kullanıcı tiplerini yükler. "
            "Form tablosu boşsa menü görünmez.",
        ),
        (
            "3. appsettings.json dosyasını yapılandırın",
            "ConnectionStrings:PostgreSqlConnection → sunucu, veritabanı, kullanıcı ve parola. "
            "Smtp bölümünü e-posta gönderimi için doldurun (Host, Port, Username, Password, FromAddress). "
            "Jwt:SecretKey en az 32 karakter olmalıdır. Hassas bilgileri üretimde ortam değişkeni veya "
            "User Secrets ile yönetin; repoya gerçek parola commit edilmemelidir.",
        ),
        (
            "4. Dosya yükleme klasörlerini oluşturun",
            "wwwroot/uploads/profiles, wwwroot/uploads/urunler ve wwwroot/uploads/kategoriler klasörlerinin "
            "var olduğundan ve uygulama kullanıcısının yazma iznine sahip olduğundan emin olun.",
        ),
        (
            "5. Bağımlılıkları yükleyin ve uygulamayı başlatın",
            "Geliştirme: `dotnet restore` ardından `dotnet run --project QRDestekliStokVeBarkodYonetimi`. "
            "Üretim: publish klasöründe `dotnet QRDestekliStokVeBarkodYonetimi.dll` veya IIS üzerinden çalıştırın. "
            "ASPNETCORE_ENVIRONMENT=Production ayarlayın.",
        ),
        (
            "6. İlk çalıştırmada admin kullanıcısı oluşur",
            "Uygulama ilk açılışta SeedAdminAsync çalıştırır. admin@sistem.com e-postası yoksa otomatik oluşturulur: "
            "E-posta: admin@sistem.com | Parola: Admin@123. Mevcut admin varsa yalnızca eksik form yetkileri tamamlanır.",
        ),
        (
            "7. İlk giriş sonrası yapılacaklar",
            "/login adresinden admin ile giriş yapın. Varsayılan admin parolasını değiştirin. "
            "Menü ve sayfa erişimlerini kontrol edin. Birim ve kategori tanımlarını girin. "
            "SMTP ayarlıysa şifre sıfırlama ve hesap onay e-postasını test edin. "
            "Diğer kullanıcılar için kullanıcı tipi ve yetki atamalarını yapın.",
        ),
        (
            "8. Üretim kontrol listesi",
            "HTTPS zorunlu kılın. Varsayılan admin parolasını değiştirin. appsettings içindeki gerçek şifreleri kaldırın. "
            "Firewall’da yalnızca gerekli portları açın. PostgreSQL erişimini uygulama sunucusu ile sınırlayın. "
            "wwwroot/uploads için düzenli yedekleme planlayın.",
        ),
    ]
    for title, detail in steps:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        add_para(doc, detail)

    add_heading(doc, "6.3 Örnek Bağlantı Ayarı", 2)
    add_para(doc, "appsettings.json — ConnectionStrings (örnek, değerler ortama göre değiştirilir):")
    add_para(
        doc,
        '"PostgreSqlConnection": "server=SUNUCU;port=5432;database=StokVeBarkodYonetimiDB;'
        'Username=KULLANICI;password=PAROLA"',
    )

    add_heading(doc, "6.4 SQL Script Dosyaları", 2)
    add_table(
        doc,
        ["Dosya", "İçerik"],
        [
            ["Scripts/sql/00_kurulum.sql", "01 + 02 scriptlerini sırayla çalıştıran ana kurulum dosyası"],
            ["Scripts/sql/01_create_tables.sql", "9 tablonun CREATE TABLE ifadeleri, indeksler ve FK ilişkileri"],
            ["Scripts/sql/02_seed_data.sql", "Form menü kayıtları, Admin/Kullanıcı tipleri, örnek birimler"],
        ],
    )
    add_para(doc, "psql örnek komutlar (proje kökünden):")
    add_para(
        doc,
        "psql -U postgres -d StokVeBarkodYonetimiDB -f Scripts/sql/01_create_tables.sql\n"
        "psql -U postgres -d StokVeBarkodYonetimiDB -f Scripts/sql/02_seed_data.sql",
    )
    add_para(
        doc,
        "Scriptler IF NOT EXISTS kullanır; mevcut veritabanına tekrar çalıştırıldığında tabloları silmez. "
        "Tam sıfırlama gerekiyorsa önce ilgili tablolar DROP edilmelidir.",
    )


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    os.makedirs(ROOT / "docs", exist_ok=True)
    doc = Document()
    set_default_font(doc)
    build_cover(doc)
    doc.add_page_break()
    build_content(doc)
    out = OUTPUT
    try:
        doc.save(str(out))
    except PermissionError:
        out = OUTPUT.with_name(OUTPUT.stem + f"_v{VERSION}" + OUTPUT.suffix)
        doc.save(str(out))
        print(f"Uyarı: Ana dosya açık; alternatif kaydedildi: {out}")
    print(f"Dokümantasyon oluşturuldu: {out}")
    print(f"Boyut: {out.stat().st_size:,} bayt")


if __name__ == "__main__":
    main()
